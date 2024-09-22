"""
This module contains classes and functions for loading and representing text documents.
"""

from docx import Document as DocxDocument

# pylint: disable=too-few-public-methods
class SimpleTextDocument:
    """
    A class to represent a text document.

    Attributes:
        page_content (str): The content of the document.
        metadata (dict): Optional metadata associated with the document.
    """
    def __init__(self, text, metadata=None):
        """
        Initializes a SimpleTextDocument instance.

        Parameters:
            text (str): The content of the document.
            metadata (dict, optional): Optional metadata associated with the document. 
            Defaults to None.
        """
        self.page_content = text
        self.metadata = metadata or {}


class DocxLoader:
    """
    A loader class for .docx files.
    """
    def __init__(self, file_path):
        """
        Initializes a DocxLoader instance.

        Parameters:
            file_path (str): The path to the .docx file to be loaded.
        """
        self.file_path = file_path

    def load(self):
        """
        Loads the .docx file and extracts text.

        Returns:
            list: A list containing a single SimpleTextDocument instance with the 
            content of the .docx file.
        """
        doc = DocxDocument(self.file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        full_text = "\n".join(text)
        return [SimpleTextDocument(full_text)]

class SimpleTextLoader:
    """
    A class to load text documents from a file.

    Attributes:
        file_path (str): The path to the text file.
    """
    def __init__(self, file_path):
        """
        Initializes a SimpleTextLoader instance.

        Parameters:
            file_path (str): The path to the text file to be loaded.
        """
        self.file_path = file_path

    def load(self):
        """
        Loads the text document from the file.

        Returns:
            list: A list containing a single SimpleTextDocument instance with the 
            content of the file.
        """
        with open(self.file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return [SimpleTextDocument(text)]
